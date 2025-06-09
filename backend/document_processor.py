import os
import PyPDF2
from docx import Document
from typing import Dict, Optional
from cache_manager import CacheManager
from logger_config import doc_logger

class DocumentProcessor:
    """文档处理器，用于解析PDF和Word文档内容"""
    
    def __init__(self):
        self.cache_manager = CacheManager()
    
    @staticmethod
    def extract_pdf_content(file_path: str) -> Optional[str]:
        """提取PDF文档内容
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            doc_logger.info(f"开始提取PDF文件内容: {file_path}")
            with open(file_path, 'rb') as file:
                doc_logger.debug("文件成功打开")
                # 创建PDF阅读器对象
                pdf_reader = PyPDF2.PdfReader(file)
                doc_logger.debug("PDF阅读器对象创建成功")
                doc_logger.info(f"PDF文件页数: {len(pdf_reader.pages)}")
                
                # 提取所有页面的文本
                text_content = []
                doc_logger.debug("开始遍历PDF页面")
                for i, page in enumerate(pdf_reader.pages):
                    doc_logger.debug(f"正在处理第 {i+1} 页")
                    text_content.append(page.extract_text())
                doc_logger.debug("所有页面文本提取完成")
                
                # 合并所有页面的文本
                result = '\n'.join(text_content)
                doc_logger.debug("页面文本合并完成")
                doc_logger.info(f"PDF内容提取成功，文本长度: {len(result)}")
                return result
        except Exception as e:
            doc_logger.error(f"提取PDF内容时出错: {type(e).__name__} - {e}", exc_info=True)
            return None
    
    @staticmethod
    def extract_docx_content(file_path: str) -> Optional[str]:
        """提取Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            doc_logger.info(f"开始提取Word文档内容: {file_path}")
            # 打开Word文档
            doc = Document(file_path)
            
            # 提取所有段落的文本
            text_content = []
            doc_logger.info(f"Word文档段落数: {len(doc.paragraphs)}")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    doc_logger.debug(f"处理第 {i+1} 个段落")
                    text_content.append(para.text)
            
            # 合并所有段落的文本
            result = '\n'.join(text_content)
            doc_logger.info(f"Word文档内容提取成功，文本长度: {len(result)}")
            return result
        except Exception as e:
            doc_logger.error(f"提取Word文档内容时出错: {type(e).__name__} - {e}", exc_info=True)
            return None
    
    def process_document(self, file_path: str, document_id: int) -> Optional[str]:
        """处理文档，根据文件类型选择相应的解析方法
        
        Args:
            file_path: 文档文件路径
            document_id: 文档ID
            
        Returns:
            str: 提取的文本内容
        """
        try:
            doc_logger.info(f"开始处理文档 {document_id}，文件路径：{file_path}")
            doc_logger.debug(f"当前工作目录：{os.getcwd()}")
            
            # 检查文件是否存在
            if not os.path.exists(file_path):
                doc_logger.error(f"错误：文件不存在 - {file_path}")
                doc_logger.debug(f"尝试访问的完整路径：{os.path.abspath(file_path)}")
                return None
                
            # 获取文件扩展名
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            doc_logger.info(f"文件类型：{ext}，文件大小：{os.path.getsize(file_path)} 字节")
            
            # 检查缓存中是否存在处理结果
            doc_logger.info(f"检查文档 {document_id} 的缓存")
            cached_data = self.cache_manager.get_from_cache(document_id)
            if cached_data:
                doc_logger.info(f"从缓存中获取文档 {document_id} 的处理结果")
                return cached_data['processed_text']
            
            # 根据文件类型选择处理方法
            processed_text = None
            doc_logger.info(f"开始提取文档内容，使用处理器：{self.__class__.__name__}")
            if ext == '.pdf':
                doc_logger.debug(f"调用 extract_pdf_content，文件路径: {file_path}")
                processed_text = DocumentProcessor.extract_pdf_content(file_path)
            elif ext in ['.docx', '.doc']:
                doc_logger.debug(f"调用 extract_docx_content，文件路径: {file_path}")
                processed_text = DocumentProcessor.extract_docx_content(file_path)
            else:
                doc_logger.error(f"不支持的文件类型: {ext}")
                return None
            
            if processed_text is None:
                doc_logger.error(f"文档 {document_id} 内容提取失败")
                return None
                
            doc_logger.info(f"文档内容提取成功，文本长度：{len(processed_text)}")
            doc_logger.debug(f"提取的文本前100个字符：{processed_text[:100]}...")
            
            # 如果处理成功，保存到缓存
            doc_logger.info(f"开始保存文档 {document_id} 的处理结果到缓存")
            cache_file_path = self.cache_manager.save_to_cache(
                document_id=document_id,
                original_file_path=file_path,
                processed_text=processed_text
            )
            if cache_file_path:
                doc_logger.info(f"文档 {document_id} 的处理结果已保存到缓存: {cache_file_path}")
            else:
                doc_logger.warning(f"警告：文档 {document_id} 的处理结果缓存保存失败")
            
            return processed_text
        except Exception as e:
            doc_logger.error(f"处理文档时发生错误: {str(e)}")
            return None